import re
from io import BytesIO

from docx import Document
from pypdf import PdfReader


PDF_CONTENT_TYPE = "application/pdf"

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


class ResumeExtractionError(Exception):
    """
    Raised when usable résumé text cannot be extracted.
    """


def normalize_resume_text(text: str) -> str:
    """
    Remove problematic characters and excessive whitespace
    while preserving meaningful line breaks.
    """

    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Replace repeated spaces and tabs with one space.
    text = re.sub(
        r"[ \t]+",
        " ",
        text,
    )

    # Remove spaces immediately before line breaks.
    text = re.sub(
        r" +\n",
        "\n",
        text,
    )

    # Limit consecutive blank lines.
    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text,
    )

    return text.strip()

def extract_pdf_text(
    file_buffer: BytesIO,
) -> str:
    """
    Extract text from a text-based PDF résumé.
    """

    try:
        reader = PdfReader(file_buffer)

    except Exception as error:
        raise ResumeExtractionError(
            "The PDF could not be opened."
        ) from error

    if reader.is_encrypted:
        try:
            decrypt_result = reader.decrypt("")

        except Exception as error:
            raise ResumeExtractionError(
                "Password-protected PDFs are not supported."
            ) from error

        if decrypt_result == 0:
            raise ResumeExtractionError(
                "Password-protected PDFs are not supported."
            )

    pages: list[str] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            page_text = page.extract_text() or ""

        except Exception as error:
            raise ResumeExtractionError(
                (
                    "Text could not be extracted from "
                    f"page {page_number}."
                )
            ) from error

        if page_text.strip():
            pages.append(page_text)

    return normalize_resume_text(
        "\n\n".join(pages)
    )



def extract_docx_text(
    file_buffer: BytesIO,
) -> str:
    """
    Extract text from DOCX paragraphs and tables.
    """

    try:
        document = Document(file_buffer)

    except Exception as error:
        raise ResumeExtractionError(
            "The DOCX file could not be opened."
        ) from error

    content: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            content.append(paragraph_text)

    # Many résumé templates use tables for layout.
    for table in document.tables:
        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                cell_text = normalize_resume_text(
                    cell.text
                )

                if cell_text:
                    cells.append(cell_text)

            if cells:
                content.append(
                    " | ".join(cells)
                )

    return normalize_resume_text(
        "\n".join(content)
    )


def extract_resume_text(
    *,
    file_buffer: BytesIO,
    content_type: str,
) -> str:
    """
    Choose the correct extraction method using the stored
    résumé content type.
    """

    if content_type == PDF_CONTENT_TYPE:
        extracted_text = extract_pdf_text(
            file_buffer
        )

    elif content_type == DOCX_CONTENT_TYPE:
        extracted_text = extract_docx_text(
            file_buffer
        )

    else:
        raise ResumeExtractionError(
            "This résumé file type is not supported."
        )

    if len(extracted_text) < 30:
        raise ResumeExtractionError(
            (
                "Very little text could be extracted. "
                "The document may be scanned, image-based, "
                "empty, or improperly formatted."
            )
        )

    return extracted_text